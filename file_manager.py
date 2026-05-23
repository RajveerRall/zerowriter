import os
import time
import traceback
from docx import Document

class FileManager:
    def __init__(self, writings_dir, draft_path):
        self.writings_dir = writings_dir
        self.draft_path = draft_path
        os.makedirs(self.writings_dir, exist_ok=True)
        print(f"[FileManager] Initialized. Writings directory: {self.writings_dir}")
        
    def has_docx_support(self):
        # Always True as we strictly require docx now
        return True

    def list_files(self):
        try:
            # Strictly list .docx files
            files = [f for f in os.listdir(self.writings_dir) if f.endswith(".docx") and not f.startswith(".")]
            files.sort(key=lambda x: os.path.getmtime(os.path.join(self.writings_dir, x)), reverse=True)
            print(f"[FileManager] Found {len(files)} .docx files in writings directory.")
            return files
        except Exception as e:
            print(f"[FileManager] Error scanning writings directory: {e}")
            return []

    def has_crash_draft(self):
        return os.path.exists(self.draft_path) and os.path.getsize(self.draft_path) > 0

    def read_crash_draft(self):
        try:
            with open(self.draft_path, 'r') as f:
                content = f.read()
                print(f"[FileManager] Loaded crash draft ({len(content)} characters).")
                return content
        except Exception as e:
            print(f"[FileManager] Error reading crash draft: {e}")
            return ""

    def clear_crash_draft(self):
        if os.path.exists(self.draft_path):
            try:
                os.remove(self.draft_path)
                print("[FileManager] Cleared crash draft.")
            except Exception as e:
                print(f"[FileManager] Error clearing crash draft: {e}")

    def save_crash_draft(self, text):
        try:
            with open(self.draft_path, 'w') as f:
                f.write(text)
        except Exception as e:
            print(f"[FileManager] Error saving crash draft: {e}")

    def load_file(self, filename):
        path = os.path.join(self.writings_dir, filename)
        print(f"[FileManager] Loading file: {path}")
        try:
            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs]
            content = "\n".join(paragraphs)
            print(f"[FileManager] Successfully loaded {filename} ({len(paragraphs)} paragraphs, {len(content)} characters).")
            return content
        except Exception as e:
            print(f"[FileManager] Error loading file {filename}: {e}")
            traceback.print_exc()
            return ""

    def save_file(self, filename, text):
        path = os.path.join(self.writings_dir, filename)
        print(f"[FileManager] Saving file to: {path}")
        try:
            doc = Document()
            lines = text.split('\n')
            for line in lines:
                doc.add_paragraph(line)
            doc.save(path)
            print(f"[FileManager] Successfully saved {filename} ({len(lines)} paragraphs, {len(text)} characters).")
            self.clear_crash_draft()
            return True
        except Exception as e:
            print(f"[FileManager] Error saving file {filename}: {e}")
            traceback.print_exc()
            return False

    def generate_filename(self):
        return f"draft_{time.strftime('%Y%m%d_%H%M%S')}.docx"
